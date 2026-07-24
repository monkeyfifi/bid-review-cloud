#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""extract_text.py — Windows原生PDF/Word取数（无需pdftotext/poppler）

输出：
  <outdir>/<stem>.lines.txt    每行: 行号<TAB>文本
  <outdir>/<stem>.tables.json  结构化表格（仅docx）

支持: .docx (python-docx) / .pdf (pypdf，Windows无需安装任何系统工具)
"""

import sys
import json
import argparse
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


def extract_docx(path):
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    doc = Document(str(path))
    lines, tables = [], []
    tbl_id = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            txt = Paragraph(child, doc).text.strip()
            if txt:
                lines.append(txt)
        elif child.tag == qn("w:tbl"):
            tbl_id += 1
            t = Table(child, doc)
            start = len(lines) + 1
            rows = []
            for row in t.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append(cells)
                lines.append("[T%d] %s" % (tbl_id, " | ".join(cells)))
            tables.append({
                "table_id": tbl_id,
                "line_start": start,
                "line_end": len(lines),
                "n_rows": len(rows),
                "n_cols": max((len(r) for r in rows), default=0),
                "rows": rows,
            })
    return lines, tables


def extract_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
    lines = [ln.rstrip() for ln in text.splitlines()]
    return lines, []


def main():
    ap = argparse.ArgumentParser(description="Windows原生PDF/Word取数")
    ap.add_argument("file")
    ap.add_argument("--outdir", default=None)
    ap.add_argument("--name", default=None)
    args = ap.parse_args()

    path = Path(args.file)
    if not path.exists():
        print("ERROR file_not_found:", path)
        sys.exit(1)

    ext = path.suffix.lower()
    if ext == ".docx":
        lines, tables = extract_docx(path)
    elif ext == ".pdf":
        lines, tables = extract_pdf(path)
    elif ext == ".doc":
        print("ERROR doc_unsupported: .doc(老Word)暂不支持，请另存为.docx或.pdf后重试")
        sys.exit(2)
    else:
        print("ERROR bad_format:", ext)
        sys.exit(2)

    outdir = Path(args.outdir) if args.outdir else path.parent / "_extracted"
    outdir.mkdir(parents=True, exist_ok=True)
    stem = args.name if args.name else path.stem

    nonempty = sum(1 for ln in lines if ln.strip())
    chars = sum(len(ln) for ln in lines)

    if nonempty == 0 or chars < 50:
        print("ERROR empty_or_low_extraction format=%s nonempty=%d chars=%d" % (ext, nonempty, chars))
        print("  提取内容为空或过少——常见原因：扫描件/图片型PDF(需OCR)、加密PDF")
        sys.exit(3)

    lines_path = outdir / (stem + ".lines.txt")
    with open(lines_path, "w", encoding="utf-8") as f:
        for i, ln in enumerate(lines, 1):
            f.write("%d\t%s\n" % (i, ln))

    tables_path = None
    if tables:
        tables_path = outdir / (stem + ".tables.json")
        with open(tables_path, "w", encoding="utf-8") as f:
            json.dump(tables, f, ensure_ascii=False, indent=1)

    print("OK format=%s" % ext)
    print("lines_total=%d lines_nonempty=%d chars=%d tables=%d" % (len(lines), nonempty, chars, len(tables)))
    print("out_lines=%s" % lines_path)
    if tables_path:
        print("out_tables=%s" % tables_path)


if __name__ == "__main__":
    main()
